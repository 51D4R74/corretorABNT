"""
Editor Científico ABNT - Processador de Documentos Acadêmicos
Extrai texto de Word/PDF preservando formatação e corrige citações/referências ABNT
"""

import re
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Set
import argparse

# Bibliotecas para processamento de documentos
try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    print("Instalando python-docx...")
    os.system("pip install python-docx")
    from docx import Document

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    print("Instalando PyPDF2...")
    os.system("pip install PyPDF2")
    import PyPDF2
    from PyPDF2 import PdfReader

try:
    import pdfplumber
except ImportError:
    print("Instalando pdfplumber...")
    os.system("pip install pdfplumber")
    import pdfplumber

try:
    from markdown import markdown
except ImportError:
    print("Instalando markdown...")
    os.system("pip install markdown")
    from markdown import markdown


class DocumentProcessor:
    """Processa documentos Word e PDF extraindo texto e formatação"""
    
    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.extension = self.file_path.suffix.lower()
        self.content = ""
        self.references = {}
        
    def extract_from_word(self) -> str:
        """Extrai texto do Word preservando formatação em Markdown"""
        doc = Document(self.file_path)
        markdown_content = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                markdown_content.append("")
                continue
                
            # Detectar níveis de título
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name[-1].isdigit() else 1
                markdown_content.append(f"{'#' * level} {para.text}")
            else:
                # Processar formatação inline (negrito, itálico)
                formatted_text = self._process_run_formatting(para)
                markdown_content.append(formatted_text)
        
        # Processar tabelas
        for table in doc.tables:
            markdown_content.append(self._convert_table_to_markdown(table))
        
        return "\n\n".join(markdown_content)
    
    def _process_run_formatting(self, paragraph) -> str:
        """Processa formatação de runs (negrito, itálico) em um parágrafo"""
        result = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
                
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            result.append(text)
        
        return "".join(result)
    
    def _convert_table_to_markdown(self, table) -> str:
        """Converte tabela do Word para Markdown"""
        md_table = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            md_table.append("| " + " | ".join(cells) + " |")
            
            # Adicionar linha separadora após o cabeçalho
            if i == 0:
                md_table.append("| " + " | ".join(["---"] * len(cells)) + " |")
        
        return "\n".join(md_table)
    
    def extract_from_pdf(self) -> str:
        """Extrai texto do PDF com preservação de formatação"""
        markdown_content = []
        
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                # Extrair texto com layout
                text = page.extract_text(layout=True)
                if text:
                    # Tentar preservar formatação básica
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            markdown_content.append(line)
        
        return "\n\n".join(markdown_content)
    
    def extract(self) -> str:
        """Extrai conteúdo baseado no tipo de arquivo"""
        if self.extension in ['.docx', '.doc']:
            self.content = self.extract_from_word()
        elif self.extension == '.pdf':
            self.content = self.extract_from_pdf()
        else:
            raise ValueError(f"Formato não suportado: {self.extension}")
        
        return self.content


class ABNTCitationProcessor:
    """Processa e corrige citações e referências conforme ABNT"""
    
    def __init__(self, content: str):
        self.content = content
        self.references = {}
        self.citations_in_text = set()
        
    def extract_references_section(self) -> Dict[str, Dict]:
        """Extrai a seção de referências do documento"""
        # Procurar seção de referências
        ref_pattern = r'(?:^|\n)(?:#+\s*)?(?:REFERÊNCIAS|Referências|REFERENCIAS|Referencias)(?:\s*\n|$)'
        match = re.search(ref_pattern, self.content, re.MULTILINE)
        
        if not match:
            print("⚠️ Seção de Referências não encontrada")
            return {}
        
        # Extrair texto após "Referências"
        ref_start = match.end()
        ref_text = self.content[ref_start:]
        
        # Parsear referências individuais
        references = {}
        
        # Padrão para identificar referências ABNT (AUTOR, ANO. Título...)
        ref_entries = re.split(r'\n\s*\n', ref_text)
        
        for entry in ref_entries:
            entry = entry.strip()
            if not entry or len(entry) < 20:
                continue
            
            # Extrair autor e ano
            author_year_pattern = r'^([A-ZÀÁÂÃÉÊÍÓÔÕÚÇ][A-ZÀÁÂÃÉÊÍÓÔÕÚÇa-zàáâãéêíóôõúç\s,]+?)(?:,|\.)\s*(\d{4})'
            match = re.search(author_year_pattern, entry)
            
            if match:
                author = match.group(1).strip()
                year = match.group(2).strip()
                
                # Extrair sobrenome principal para citação
                last_name = self._extract_last_name(author)
                key = f"{last_name.upper()}, {year}"
                
                references[key] = {
                    'full_author': author,
                    'year': year,
                    'last_name': last_name,
                    'full_entry': entry,
                    'has_url': bool(re.search(r'https?://', entry)),
                    'url': self._extract_url(entry),
                    'access_date': self._extract_access_date(entry)
                }
        
        self.references = references
        return references
    
    def _extract_last_name(self, author_string: str) -> str:
        """Extrai o sobrenome principal do autor"""
        # Remover "et al.", vírgulas extras
        author_string = re.sub(r'\s+et\s+al\.?', '', author_string, flags=re.IGNORECASE)
        
        # Pegar primeira parte antes da vírgula ou ponto
        parts = re.split(r'[,;.]', author_string)
        if parts:
            last_name = parts[0].strip()
            # Remover nomes intermediários comuns
            last_name = re.sub(r'\s+(de|da|do|dos|das|e)\s+', ' ', last_name, flags=re.IGNORECASE)
            return last_name.split()[-1] if ' ' in last_name else last_name
        
        return author_string.split()[0] if author_string else ""
    
    def _extract_url(self, entry: str) -> str:
        """Extrai URL da referência"""
        url_pattern = r'(?:Disponível em:|Available at:)?\s*(https?://[^\s]+)'
        match = re.search(url_pattern, entry)
        return match.group(1).strip() if match else ""
    
    def _extract_access_date(self, entry: str) -> str:
        """Extrai data de acesso da referência"""
        date_pattern = r'Acesso em:\s*(\d{1,2}\s+\w+\.?\s+\d{4})'
        match = re.search(date_pattern, entry)
        return match.group(1).strip() if match else ""
    
    def convert_numeric_to_author_date(self) -> str:
        """Converte citações numéricas [1] para formato autor-data (AUTOR, ano)"""
        content = self.content
        
        # Padrão para citações numéricas: [1], [2,3], [4-6], etc.
        numeric_pattern = r'\[(\d+(?:\s*[-,]\s*\d+)*)\]'
        
        def replace_citation(match):
            numbers = match.group(1)
            # Por enquanto, placeholder - necessário mapear números para referências
            # Em implementação real, seria necessário um mapeamento
            return f"(AUTOR, ANO)"  # Placeholder
        
        content = re.sub(numeric_pattern, replace_citation, content)
        self.content = content
        return content
    
    def add_missing_citations(self) -> str:
        """Adiciona citações faltantes para termos técnicos"""
        # Padrões de termos técnicos que necessitam citação
        technical_terms = [
            r'sinal do morcego',
            r'protocolo BLUE',
            r'sinal da praia',
            r'ultrassom point-of-care',
            r'POCUS',
        ]
        
        content = self.content
        
        for term_pattern in technical_terms:
            # Verificar se termo existe sem citação logo após
            pattern = rf'({term_pattern})(?!\s*\([A-Z][A-Za-z]+,\s*\d{{4}}\))'
            
            if re.search(pattern, content, re.IGNORECASE):
                # Buscar referência apropriada
                citation = self._find_appropriate_citation(term_pattern)
                if citation:
                    content = re.sub(
                        pattern,
                        rf'\1 {citation}',
                        content,
                        flags=re.IGNORECASE,
                        count=1
                    )
        
        self.content = content
        return content
    
    def _find_appropriate_citation(self, term: str) -> str:
        """Encontra citação apropriada para um termo técnico"""
        # Buscar nas referências por palavras-chave
        term_lower = term.lower()
        
        for key, ref in self.references.items():
            if term_lower in ref['full_entry'].lower():
                return f"({key})"
        
        return ""
    
    def format_references_abnt(self) -> str:
        """Formata lista de referências conforme ABNT NBR 6023"""
        if not self.references:
            return self.content
        
        # Localizar seção de referências
        ref_pattern = r'(?:^|\n)((?:#+\s*)?(?:REFERÊNCIAS|Referências))(?:\s*\n)'
        match = re.search(ref_pattern, self.content, re.MULTILINE)
        
        if not match:
            return self.content
        
        ref_start = match.start()
        content_before = self.content[:match.end()]
        
        # Ordenar referências alfabeticamente
        sorted_refs = sorted(self.references.items(), key=lambda x: x[1]['last_name'].upper())
        
        # Formatar cada referência
        formatted_refs = []
        current_date = datetime.now().strftime("%d %b. %Y")
        
        for key, ref in sorted_refs:
            entry = ref['full_entry']
            
            # Verificar e atualizar data de acesso se necessário
            if ref['has_url'] and ref['url']:
                # Garantir formato correto do link
                if 'Disponível em:' not in entry:
                    entry = entry.replace(ref['url'], f"Disponível em: {ref['url']}")
                
                # Atualizar data de acesso se antiga ou ausente
                if ref['access_date']:
                    # Verificar se data é antiga (implementação simplificada)
                    entry = re.sub(
                        r'Acesso em:\s*\d{1,2}\s+\w+\.?\s+\d{4}',
                        f'Acesso em: {current_date}',
                        entry
                    )
                else:
                    # Adicionar data de acesso
                    entry += f" Acesso em: {current_date}."
            
            formatted_refs.append(entry)
        
        # Reconstruir documento
        new_content = content_before + "\n\n" + "\n\n".join(formatted_refs)
        self.content = new_content
        return new_content
    
    def validate_citation_reference_match(self) -> Tuple[Set[str], Set[str]]:
        """Valida correspondência 1:1 entre citações no texto e referências"""
        # Extrair todas as citações no texto
        citation_pattern = r'\(([A-Z][A-ZÀÁÂÃÉÊÍÓÔÕÚÇa-z]+(?:\s+et\s+al\.?)?),\s*(\d{4})\)'
        citations = re.findall(citation_pattern, self.content)
        
        citations_in_text = {f"{author.upper()}, {year}" for author, year in citations}
        references_listed = set(self.references.keys())
        
        # Citações sem referência
        missing_refs = citations_in_text - references_listed
        # Referências não citadas
        unused_refs = references_listed - citations_in_text
        
        return missing_refs, unused_refs
    
    def remove_html_tags(self) -> str:
        """Remove tags HTML e substitui por caracteres Unicode apropriados"""
        content = self.content
        
        # Substituições comuns
        replacements = {
            r'<sub>(\d+)</sub>': r'₀₁₂₃₄₅₆₇₈₉'[int(r'\1')] if r'\1'.isdigit() else r'\1',
            r'<sup>(\d+)</sup>': r'⁰¹²³⁴⁵⁶⁷⁸⁹'[int(r'\1')] if r'\1'.isdigit() else r'\1',
            r'<br\s*/?>': '\n',
            r'</?p>': '\n',
            r'</?div[^>]*>': '',
            r'</?span[^>]*>': '',
        }
        
        for pattern, replacement in replacements.items():
            content = re.sub(pattern, replacement, content, flags=re.IGNORECASE)
        
        # Remover quaisquer tags HTML restantes
        content = re.sub(r'<[^>]+>', '', content)
        
        self.content = content
        return content
    
    def process(self) -> str:
        """Executa pipeline completo de processamento"""
        print("📚 Extraindo referências...")
        self.extract_references_section()
        print(f"   ✓ {len(self.references)} referências encontradas")
        
        print("🔄 Convertendo citações numéricas para autor-data...")
        self.convert_numeric_to_author_date()
        
        print("➕ Adicionando citações faltantes...")
        self.add_missing_citations()
        
        print("📋 Formatando referências ABNT...")
        self.format_references_abnt()
        
        print("✓ Validando correspondência citações/referências...")
        missing, unused = self.validate_citation_reference_match()
        if missing:
            print(f"   ⚠️ Citações sem referência: {missing}")
        if unused:
            print(f"   ⚠️ Referências não citadas: {unused}")
        
        print("🧹 Removendo tags HTML...")
        self.remove_html_tags()
        
        return self.content


class ScientificEditorABNT:
    """Editor científico completo para documentos acadêmicos"""
    
    def __init__(self, input_file: str, output_file: str = None):
        self.input_file = input_file
        self.output_file = output_file or self._generate_output_filename()
        
    def _generate_output_filename(self) -> str:
        """Gera nome do arquivo de saída"""
        path = Path(self.input_file)
        return str(path.parent / f"{path.stem}_edited.md")
    
    def process_document(self) -> str:
        """Processa documento completo"""
        print(f"\n🔬 Editor Científico ABNT")
        print(f"📄 Arquivo: {self.input_file}\n")
        
        # Etapa 1: Extrair conteúdo
        print("📖 Etapa 1: Extraindo conteúdo do documento...")
        processor = DocumentProcessor(self.input_file)
        content = processor.extract()
        print(f"   ✓ {len(content)} caracteres extraídos\n")
        
        # Etapa 2: Processar citações e referências
        print("📝 Etapa 2: Processando citações e referências ABNT...")
        citation_processor = ABNTCitationProcessor(content)
        processed_content = citation_processor.process()
        print(f"   ✓ Processamento concluído\n")
        
        # Etapa 3: Salvar resultado
        print(f"💾 Etapa 3: Salvando documento editado...")
        self._save_output(processed_content)
        print(f"   ✓ Salvo em: {self.output_file}\n")
        
        return processed_content
    
    def _save_output(self, content: str):
        """Salva conteúdo processado em arquivo Markdown"""
        with open(self.output_file, 'w', encoding='utf-8') as f:
            # Adicionar tag immersive conforme requisito
            f.write("<immersive>\n\n")
            f.write(content)
            f.write("\n\n</immersive>")
        
        print(f"✅ Documento processado com sucesso!")
        print(f"📊 Estatísticas:")
        print(f"   • Tamanho: {len(content)} caracteres")
        print(f"   • Linhas: {content.count(chr(10))} linhas")


def main():
    """Função principal"""
    parser = argparse.ArgumentParser(
        description='Editor Científico ABNT - Processa documentos Word/PDF e corrige citações/referências'
    )
    parser.add_argument('input', help='Arquivo de entrada (Word ou PDF)')
    parser.add_argument('-o', '--output', help='Arquivo de saída (Markdown)', default=None)
    
    args = parser.parse_args()
    
    # Verificar se arquivo existe
    if not os.path.exists(args.input):
        print(f"❌ Erro: Arquivo não encontrado: {args.input}")
        return
    
    # Processar documento
    editor = ScientificEditorABNT(args.input, args.output)
    editor.process_document()


if __name__ == "__main__":
    main()
