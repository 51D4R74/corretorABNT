"""
Exportador de documentos Word (DOCX) com formatação preservada
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, List


class DocxExporter:
    """
    Exporta conteúdo processado para formato Word (.docx)
    - Preserva formatação (negrito, itálico, títulos)
    - Remove tags HTML
    - Mantém estrutura de parágrafos e espaçamento
    - Formata referências em ABNT
    """
    
    def __init__(self, content: str):
        self.content = content
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Configura estilos do documento"""
        # Configurar estilo normal
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Configurar margens (2.5cm superior/esquerda, 2cm inferior/direita)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(1)
            section.right_margin = Inches(0.8)
    
    def remove_html_tags(self) -> str:
        """
        Remove todas as tags HTML antes de processar
        
        Returns:
            Texto sem HTML
        """
        content = self.content
        
        # Preservar tags de formatação temporariamente
        content = re.sub(r'<strong>(.*?)</strong>', r'**\1**', content, flags=re.IGNORECASE)
        content = re.sub(r'<b>(.*?)</b>', r'**\1**', content, flags=re.IGNORECASE)
        content = re.sub(r'<em>(.*?)</em>', r'*\1*', content, flags=re.IGNORECASE)
        content = re.sub(r'<i>(.*?)</i>', r'*\1*', content, flags=re.IGNORECASE)
        
        # Subscritos - marcar para processamento posterior
        content = re.sub(r'<sub>(.*?)</sub>', r'_{{\1}}', content, flags=re.IGNORECASE)
        
        # Sobrescritos - marcar para processamento posterior
        content = re.sub(r'<sup>(.*?)</sup>', r'^{{\1}}', content, flags=re.IGNORECASE)
        
        # Quebras de linha
        content = re.sub(r'<br\s*/?>', '\n', content, flags=re.IGNORECASE)
        
        # Parágrafos
        content = re.sub(r'<p>', '\n\n', content, flags=re.IGNORECASE)
        content = re.sub(r'</p>', '\n\n', content, flags=re.IGNORECASE)
        
        # Remover spans e divs
        content = re.sub(r'</?(?:span|div)[^>]*>', '', content, flags=re.IGNORECASE)
        
        # Remover tag immersive
        content = re.sub(r'</?immersive>', '', content, flags=re.IGNORECASE)
        
        # Remover quaisquer outras tags HTML
        content = re.sub(r'<[^>]+>', '', content)
        
        # Decodificar entidades HTML
        html_entities = {
            '&nbsp;': ' ',
            '&lt;': '<',
            '&gt;': '>',
            '&amp;': '&',
            '&quot;': '"',
            '&#39;': "'",
            '&mdash;': '—',
            '&ndash;': '–',
        }
        
        for entity, char in html_entities.items():
            content = content.replace(entity, char)
        
        self.content = content
        return content
    
    def _parse_inline_formatting(self, text: str) -> List[Dict]:
        """
        Parseia formatação inline (negrito, itálico, sub/sobrescrito)
        
        Args:
            text: Texto com marcações
            
        Returns:
            Lista de segmentos com formatação
        """
        segments = []
        
        # Padrão para encontrar formatações
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|\_{.*?\}|\^{.*?\}|[^*_^]+)'
        matches = re.finditer(pattern, text)
        
        for match in matches:
            segment = match.group(0)
            
            if not segment:
                continue
            
            # Negrito + Itálico
            if segment.startswith('***') and segment.endswith('***'):
                segments.append({
                    'text': segment[3:-3],
                    'bold': True,
                    'italic': True
                })
            # Negrito
            elif segment.startswith('**') and segment.endswith('**'):
                segments.append({
                    'text': segment[2:-2],
                    'bold': True,
                    'italic': False
                })
            # Itálico
            elif segment.startswith('*') and segment.endswith('*'):
                segments.append({
                    'text': segment[1:-1],
                    'bold': False,
                    'italic': True
                })
            # Subscrito
            elif segment.startswith('_{') and segment.endswith('}'):
                segments.append({
                    'text': segment[2:-1],
                    'subscript': True
                })
            # Sobrescrito
            elif segment.startswith('^{') and segment.endswith('}'):
                segments.append({
                    'text': segment[2:-1],
                    'superscript': True
                })
            # Texto normal
            else:
                segments.append({
                    'text': segment,
                    'bold': False,
                    'italic': False
                })
        
        return segments
    
    def _add_formatted_paragraph(self, text: str, is_heading: int = 0, is_reference: bool = False):
        """
        Adiciona parágrafo com formatação ao documento
        
        Args:
            text: Texto do parágrafo
            is_heading: Nível do título (0 = parágrafo normal)
            is_reference: Se é uma referência bibliográfica
        """
        if not text.strip():
            # Adicionar parágrafo vazio para espaçamento
            self.doc.add_paragraph()
            return
        
        # Criar parágrafo
        if is_heading > 0:
            # Títulos
            level = min(is_heading, 3)  # Limitar a 3 níveis
            para = self.doc.add_heading(level=level)
            para.clear()
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(16 - level * 2)
        else:
            para = self.doc.add_paragraph()
            
            # Justificar texto (padrão ABNT)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Parsear e adicionar segmentos com formatação
            segments = self._parse_inline_formatting(text)
            
            for segment in segments:
                run = para.add_run(segment['text'])
                
                # Aplicar formatação
                if segment.get('bold'):
                    run.bold = True
                if segment.get('italic'):
                    run.italic = True
                if segment.get('subscript'):
                    run.font.subscript = True
                if segment.get('superscript'):
                    run.font.superscript = True
                
                # Fonte padrão
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
                # Referências: autor em negrito maiúsculo
                if is_reference and segment == segments[0]:
                    run.bold = True
    
    def _process_content(self):
        """Processa o conteúdo e adiciona ao documento Word"""
        # Remover HTML primeiro
        self.remove_html_tags()
        
        # Normalizar espaços
        content = re.sub(r' +', ' ', self.content)
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # Dividir em linhas
        lines = content.split('\n')
        
        in_references = False
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # Linha vazia
            if not line:
                continue
            
            # Detectar títulos Markdown
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2)
                self._add_formatted_paragraph(title, is_heading=level)
                
                # Detectar seção de referências
                if re.search(r'referências', title, re.IGNORECASE):
                    in_references = True
                continue
            
            # Detectar referências ABNT (começa com AUTOR em maiúsculo)
            is_ref = in_references and re.match(r'^[A-ZÀÁÂÃÉÊÍÓÔÕÚÇ]', line)
            
            # Adicionar parágrafo normal ou referência
            self._add_formatted_paragraph(line, is_reference=is_ref)
    
    def export(self, output_path: str):
        """
        Exporta conteúdo para arquivo Word
        
        Args:
            output_path: Caminho do arquivo de saída (.docx)
        """
        # Processar conteúdo e construir documento
        self._process_content()
        
        # Salvar documento
        self.doc.save(output_path)
    
    def validate_document(self) -> Dict[str, bool]:
        """
        Valida se documento foi criado corretamente
        
        Returns:
            Dict com validações
        """
        validations = {
            'has_paragraphs': len(self.doc.paragraphs) > 0,
            'has_proper_font': any(
                p.runs and p.runs[0].font.name == 'Times New Roman' 
                for p in self.doc.paragraphs if p.runs
            ),
            'has_formatting': any(
                p.runs and (p.runs[0].bold or p.runs[0].italic)
                for p in self.doc.paragraphs if p.runs
            ),
        }
        
        return validations
