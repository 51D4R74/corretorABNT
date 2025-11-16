"""
Extrator de documentos Word (.docx) com preservação completa de formatação
Baseado em python-docx (enterprise-level)
"""

from typing import Dict, List, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class WordExtractor:
    """
    Extrai texto de documentos Word preservando:
    - Negritos, itálicos
    - Estrutura de parágrafos
    - Espaçamentos
    - Quebras de linha
    - Tabelas
    - Hierarquia de títulos
    """
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = Document(file_path)
        self.content_blocks = []
        
    def extract(self) -> Dict:
        """
        Extrai todo conteúdo do documento preservando estrutura
        
        Returns:
            Dict com estrutura completa do documento
        """
        result = {
            'paragraphs': [],
            'tables': [],
            'metadata': self._extract_metadata(),
            'structure': []
        }
        
        # Processar todos os parágrafos
        for para in self.document.paragraphs:
            para_data = self._extract_paragraph(para)
            if para_data:
                result['paragraphs'].append(para_data)
                result['structure'].append({
                    'type': 'paragraph',
                    'data': para_data
                })
        
        # Processar todas as tabelas
        for table in self.document.tables:
            table_data = self._extract_table(table)
            result['tables'].append(table_data)
            result['structure'].append({
                'type': 'table',
                'data': table_data
            })
        
        return result
    
    def _extract_metadata(self) -> Dict:
        """Extrai metadados do documento"""
        core_props = self.document.core_properties
        return {
            'author': core_props.author or '',
            'title': core_props.title or '',
            'subject': core_props.subject or '',
            'created': core_props.created,
            'modified': core_props.modified
        }
    
    def _extract_paragraph(self, para: Paragraph) -> Dict:
        """
        Extrai parágrafo com toda formatação preservada
        
        Args:
            para: Objeto Paragraph do python-docx
            
        Returns:
            Dict com dados do parágrafo e formatação
        """
        if not para.text.strip():
            return None
        
        para_data = {
            'text': para.text,
            'style': para.style.name if para.style else 'Normal',
            'alignment': self._get_alignment(para),
            'runs': [],
            'is_heading': para.style.name.startswith('Heading') if para.style else False,
            'heading_level': self._get_heading_level(para)
        }
        
        # Processar runs individuais para preservar formatação inline
        for run in para.runs:
            run_data = self._extract_run(run)
            if run_data:
                para_data['runs'].append(run_data)
        
        return para_data
    
    def _extract_run(self, run: Run) -> Dict:
        """
        Extrai formatação de um run (sequência de texto com formatação uniforme)
        
        Args:
            run: Objeto Run do python-docx
            
        Returns:
            Dict com texto e formatação
        """
        if not run.text:
            return None
        
        return {
            'text': run.text,
            'bold': run.bold if run.bold is not None else False,
            'italic': run.italic if run.italic is not None else False,
            'underline': run.underline if run.underline is not None else False,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size.pt if run.font.size else None,
            'color': self._get_color(run)
        }
    
    def _get_color(self, run: Run) -> str:
        """Extrai cor do texto em formato RGB"""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
        except:
            pass
        return None
    
    def _get_alignment(self, para: Paragraph) -> str:
        """Retorna alinhamento do parágrafo"""
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: 'left',
            WD_PARAGRAPH_ALIGNMENT.CENTER: 'center',
            WD_PARAGRAPH_ALIGNMENT.RIGHT: 'right',
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'justify'
        }
        return alignment_map.get(para.alignment, 'left')
    
    def _get_heading_level(self, para: Paragraph) -> int:
        """Extrai nível do título (1-6)"""
        if para.style and para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split()[-1])
                return level
            except:
                return 0
        return 0
    
    def _extract_table(self, table: Table) -> Dict:
        """
        Extrai tabela com formatação preservada
        
        Args:
            table: Objeto Table do python-docx
            
        Returns:
            Dict com estrutura da tabela
        """
        table_data = {
            'rows': [],
            'row_count': len(table.rows),
            'col_count': len(table.columns) if table.rows else 0
        }
        
        for row in table.rows:
            row_data = {
                'cells': []
            }
            
            for cell in row.cells:
                cell_data = {
                    'text': cell.text.strip(),
                    'paragraphs': []
                }
                
                # Extrair parágrafos dentro da célula
                for para in cell.paragraphs:
                    para_data = self._extract_paragraph(para)
                    if para_data:
                        cell_data['paragraphs'].append(para_data)
                
                row_data['cells'].append(cell_data)
            
            table_data['rows'].append(row_data)
        
        return table_data
    
    def to_markdown(self) -> str:
        """
        Converte documento extraído para Markdown preservando formatação
        
        Returns:
            String Markdown com formatação completa
        """
        data = self.extract()
        markdown_lines = []
        
        for item in data['structure']:
            if item['type'] == 'paragraph':
                md = self._paragraph_to_markdown(item['data'])
                if md:
                    markdown_lines.append(md)
            elif item['type'] == 'table':
                md = self._table_to_markdown(item['data'])
                if md:
                    markdown_lines.append(md)
        
        return '\n\n'.join(markdown_lines)
    
    def _paragraph_to_markdown(self, para_data: Dict) -> str:
        """Converte dados do parágrafo para Markdown"""
        # Títulos
        if para_data['is_heading'] and para_data['heading_level']:
            level = para_data['heading_level']
            return f"{'#' * level} {para_data['text']}"
        
        # Parágrafo normal com formatação inline
        text_parts = []
        for run in para_data['runs']:
            text = run['text']
            
            # Aplicar formatação
            if run['bold'] and run['italic']:
                text = f"***{text}***"
            elif run['bold']:
                text = f"**{text}**"
            elif run['italic']:
                text = f"*{text}*"
            
            text_parts.append(text)
        
        return ''.join(text_parts) if text_parts else para_data['text']
    
    def _table_to_markdown(self, table_data: Dict) -> str:
        """Converte tabela para formato Markdown"""
        markdown_rows = []
        
        for i, row in enumerate(table_data['rows']):
            cells = [cell['text'] for cell in row['cells']]
            markdown_rows.append('| ' + ' | '.join(cells) + ' |')
            
            # Linha separadora após primeira linha (cabeçalho)
            if i == 0:
                separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                markdown_rows.append(separator)
        
        return '\n'.join(markdown_rows)
